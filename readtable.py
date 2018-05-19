# -*- coding: utf-8 -*-
# -*- coding: utf-8 -*-
"""
Spyder Editor

This is a temporary script file.
"""
import os
import win32com
from win32com.client import Dispatch, constants
from docx import Document
from docx.shared import Inches
import re
import pandas as pd
import xlwt


#path= "D:\\wjh\\"
#file = "2017-07-17-尚品宅配：2017年7月14日投资者关系活动记录表.DOCX"

column0 = [u'文件名称',u'股票代码',u'股票名称',u'时间头',u'时间',u'参与单位与人员头',u'参与单位与人员头',u'活动类别头',u'活动类别']   

def findactty(desc):
    res = ''
    for i in desc.split():
        if i.find("■") !=-1 or i.find("√") !=-1  or i.find("☑") !=-1:
            res = res+" "+i[1:len(i)]
    return res

def splitname(names):
    names="广发证券股份有限公司：彭雾；中信证券股份有限公司：贾常涛。"
    ss = re.split('[：\s；。——、）（]',names)
    print(ss)
    companys = []
    persons = []
    curcompany ="空公司"
    for i in ss:
        if len(i)>3:
            companys.append(i)
            curcompany = i
        if 1<len(i)<=3:
            persons.append(curcompany+"%"+i)
    print("全部AAAAAAAAAAAAAAAAAAA："+names)
    print(companys)
    print("""""""""""""""""")
    print(persons)
    return companys,persons

def write_excel(df):  
    f = xlwt.Workbook()
    sheet1 = f.add_sheet(u'sheet1',cell_overwrite_ok=True) #创建sheet  
     
    for i in range(0,len(column0)):  
        sheet1.write(0,i,column0[i])

    for j in range(0,len(df)):
        sheet1.write(j+1,0,  df[j]["file"])
        sheet1.write(j+1,1,  df[j]["code"])
        sheet1.write(j+1,2,  df[j]["name"])
        #sheet1.write(j+1,3,  df[j]["time_head"])
        sheet1.write(j+1,4,  df[j]["time_des"])
        #sheet1.write(j+1,5,  df[j]["group_head"])
        sheet1.write(j+1,6,  df[j]["group_dec"])
        #sheet1.write(j+1,7,  df[j]["act_head"])
        sheet1.write(j+1,8,  findactty(df[j]["act_des"]))
    f.save('demo1.xlsx')

w = win32com.client.Dispatch('Word.Application')
def parse_doc(path,file):
    doc = w.Documents.Open( FileName = path+file )
    t = doc.Tables[0]
    p1 = doc.Paragraphs[0]
    ss = re.split('[：\s]',str(p1))
    ss = [x for x in ss if len(x)>0]
    code = ss[1]
    name = ss[-1]
    act_head = t.Rows[0].Cells[0].Range.Text
    act_des = t.Rows[0].Cells[1].Range.Text
    group_head = t.Rows[1].Cells[0].Range.Text
    group_dec = t.Rows[1].Cells[1].Range.Text
    time_head = t.Rows[2].Cells[0].Range.Text 
    time_des = t.Rows[2].Cells[1].Range.Text 
    doc.Close()
    return {"file":file,"code":code,"name":name,"time_head":time_head,"time_des":time_des,"group_head":group_head,"group_dec":group_dec,"act_head":act_head,"act_des":act_des}

def parse_docx(path,file):
    document = Document(path+file)    
    l = [ paragraph.text for paragraph in document.paragraphs];
    ss = re.split('[：\s]',l[0])
    code = ss[1]
    name = ss[-1]
    table = document.tables[0]
    hdr_cells = table.rows[0].cells
    act_head = hdr_cells[0].text
    act_des = hdr_cells[1].text
    hdr_cells = table.rows[1].cells
    group_head = hdr_cells[0].text
    group_dec = hdr_cells[1].text
    #splitname(group_dec)
    hdr_cells = table.rows[2].cells
    time_head = hdr_cells[0].text 
    time_des = hdr_cells[1].text 
    return {"file":file,"code":code,"name":name,"time_head":time_head,"time_des":time_des,"group_head":group_head,"group_dec":group_dec,"act_head":act_head,"act_des":act_des}

def parse_pdf(path,file):
    document = Document(path+file)    
    l = [ paragraph.text for paragraph in document.paragraphs];
    ss = re.split('[：\s]',l[0])
    code = ss[1]
    name = ss[-1]
    table = document.tables[0]
    hdr_cells = table.rows[0].cells
    act_head = hdr_cells[0].text
    act_des = hdr_cells[1].text
    hdr_cells = table.rows[1].cells
    group_head = hdr_cells[0].text
    group_dec = hdr_cells[1].text
    #splitname(group_dec)
    hdr_cells = table.rows[2].cells
    time_head = hdr_cells[0].text 
    time_des = hdr_cells[1].text 
    return {"file":file,"code":code,"name":name,"time_head":time_head,"time_des":time_des,"group_head":group_head,"group_dec":group_dec,"act_head":act_head,"act_des":act_des}

if __name__ == "__main__":
    PATH = "D:\\wjh\\" 
    s = []
    doc_files = os.listdir(PATH)
    for doc in doc_files[0:10]:
        if (os.path.splitext(doc)[1] == '.DOCX' or os.path.splitext(doc)[1] == '.docx') and str(doc).find("投资者关系活动记录表") !=-1:
            try:
                res = parse_docx(PATH,doc)
                #print(res)
                s.append(res)
                #print(len(s))
            except Exception as e:
                print(e)
        elif (os.path.splitext(doc)[1] == '.DOC' or os.path.splitext(doc)[1] == '.doc') and str(doc).find("投资者关系活动记录表") !=-1:
            try:
                res = parse_doc(PATH,doc)
                print(res)
                s.append(res)
                #print(len(s))
            except Exception as e:
                print(e)
        elif (os.path.splitext(doc)[1] == '.PDF' or os.path.splitext(doc)[1] == '.pdf') and str(doc).find("投资者关系活动记录表") !=-1:
            try:
                res = parse_pdf(PATH,doc)
                print(res)
                s.append(res)
                #print(len(s))
            except Exception as e:
                print(e)
    write_excel(s)