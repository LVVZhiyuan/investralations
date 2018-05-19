# -*- coding: utf-8 -*-
# -*- coding: utf-8 -*-
"""
Spyder Editor

This is a temporary script file.
"""
from docx import Document
from docx.shared import Inches
import re

def findty(desc):
    res = ''
    for i in desc.split():
        if i.find("■") !=-1 or i.find("√") !=-1  or i.find("☑") !=-1:
            res = res+" "+i[1:len(i)]
    return res

def findname(names):
    ss = re.split('[：\s；。——、）（]',names)
    companys = []
    persons = []
    curcompany ="空公司"
    for i in ss:
        if len(i)>3:
            companys.append(i)
            curcompany = i
        if 1<len(i)<=3:
            persons.append(curcompany+"%"+i)
    return persons


def parse_docx(path,file):
    document = Document(path+file)    
    l = [ paragraph.text for paragraph in document.paragraphs]
    for i in l:
        if i.count('证券代码')>0:        
            ss = re.split('[：\s]',i)
            code = ss[1]
            name = ss[-1]
    table = document.tables[0]
    hdr_cells = table.rows[0].cells
    ty_head = hdr_cells[0].text
    pre_ty = hdr_cells[1].text
    hdr_cells = table.rows[1].cells
    pss_head = hdr_cells[0].text
    pss = hdr_cells[1].text
    hdr_cells = table.rows[2].cells
    tme_head = hdr_cells[0].text 
    tme = hdr_cells[1].text 
    
    master = table.rows[4].cells[1].text
    content = table.rows[5].cells[1].text
    content_len = str(len(content))

    asks = max(content.count('？'),content.count('答：'))
    asks = max(asks,content.count('问：'))
    
    ty = findty(pre_ty)
    psslist = findname(pss)
    return {"file":file,"code":code,"name":name,"tme_head":tme_head,"tme":tme,
            "pss_head":pss_head,"pss":pss,"ty_head":ty_head,"ty":ty,"pre_ty":pre_ty,'psslist':psslist,'master':master,'content_len':content_len,'asks':asks}
    #return {"file":file,"code":code,"name":name,"tme_head":tme_head,"tme":tme,"pss_head":pss_head,"pss":pss,"ty_head":ty_head,"ty":ty,"pre_ty":pre_ty,'psslist':psslist}

if __name__ == "__main__":
    path = u'D:\\wjh\\2015\\'
    file = u'2015-01-19-东华软件：2015年1月16日投资者关系活动记录表.DOCX'
    extractdic = parse_docx(path,file)
    for i in extractdic:
        print(i,':',extractdic[i],'\n')